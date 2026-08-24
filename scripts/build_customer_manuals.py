from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "客户培训手册"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GRAY = "666666"
RED = "9B1C1C"


def font(run, size=11, bold=False, color="000000", name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            margins(cell)


def setup(doc, short_title):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Title", 30, 0, 10, NAVY), ("Subtitle", 14, 0, 8, GRAY),
        ("Heading 1", 16, 18, 10, BLUE), ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 11.5, 10, 5, NAVY)):
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = name != "Subtitle"
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run(f"企业知识智能平台  |  {short_title}"), 8.5, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("商业交付培训资料  ·  2026-08-24"), 8, color=GRAY)


def cover(doc, title, subtitle, audience):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("ENTERPRISE KNOWLEDGE INTELLIGENCE"), 13, True, BLUE)
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle, style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(45)
    info = table(doc, [("适用对象", audience), ("产品版本", "客户交付版（以现场安装包为准）"),
                       ("文档版本", "V1.1"), ("保密等级", "客户内部资料"),
                       ("编制日期", "2026 年 8 月 24 日"), ("客户环境", "____________________________")],
                 [1900, 7460], header=False)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    note(doc, "说明", "本文档依据当前项目仓库与界面编写。客户现场的域名、账号、模型供应商和启用模块应以交付配置为准。")
    doc.add_page_break()


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), 10.5, True)
        font(p.add_run(text[len(bold_prefix):]), 10.5)
    else:
        font(p.add_run(text), 10.5)
    return p


def bullets(doc, items, numbered=False):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        font(p.add_run("•  "), 10.5, True, color=NAVY)
        font(p.add_run(item), 10.5)


def note(doc, label, text, warning=False):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [9360])
    shade(t.cell(0, 0), "FDECEC" if warning else PALE)
    p = t.cell(0, 0).paragraphs[0]
    font(p.add_run(label + "："), 10, True, RED if warning else NAVY)
    font(p.add_run(text), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def table(doc, rows, widths, headers=None, header=True):
    data = ([headers] if headers else []) + [list(r) for r in rows]
    t = doc.add_table(rows=len(data), cols=len(widths))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    for r, values in enumerate(data):
        for c, value in enumerate(values):
            cell = t.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if (headers and r == 0) or (header and c == 0):
                shade(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), 9.5, bold=((headers and r == 0) or (header and c == 0)))
    return t


def add_image(doc, filename, caption, width=6.2):
    path = ROOT / "docs" / "images" / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        font(c.add_run(caption), 8.5, color=GRAY)


def user_manual():
    doc = Document()
    setup(doc, "用户使用手册")
    cover(doc, "用户使用手册", "客户培训版 · 从知识入库到可信问答", "业务用户、知识库管理员、租户管理员、平台管理员")
    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_before = Pt(0)
    toc_title.paragraph_format.space_after = Pt(12)
    font(toc_title.add_run("目录"), 20, True, NAVY)
    toc_placeholder = doc.add_paragraph("[[TOC]]")
    toc_placeholder.paragraph_format.space_after = Pt(0)
    doc.add_page_break()
    heading(doc, "1. 手册定位与培训目标")
    para(doc, "完成本手册后，学员应能独立登录系统、创建并维护知识库、完成文档入库与检索问答，并能在权限范围内使用智能体。")
    table(doc, [("平台管理员", "管理租户、用户与平台级智能体，可切换全部租户"),
                ("租户管理员", "管理本租户用户、知识库、租户信息、审核与评测"),
                ("普通成员", "使用被授权的知识库、搜索和问答能力")], [2100, 7260], headers=["角色", "主要职责"])

    heading(doc, "2. 首次登录与安全检查")
    bullets(doc, ["在浏览器打开客户交付地址（示例：http://服务器地址:前端端口）。", "输入实施人员或管理员通过受控渠道提供的初始用户名和密码。", "首次登录后立即修改初始密码，并确认当前租户和角色正确。", "平台管理员日常业务操作前，先切换到目标租户，避免在错误租户中配置。"], numbered=True)
    note(doc, "安全提醒", "默认密码仅用于首次初始化。生产环境不得长期使用默认密码，也不要通过聊天、邮件明文传递 API Key。", True)
    heading(doc, "2.1 个人账号操作", 2)
    bullets(doc, ["在用户菜单中确认当前用户名、角色和租户。", "如界面提供密码修改入口，按密码策略设置新密码；如未提供自助入口，联系租户管理员在用户管理中重置。", "使用完成后从用户菜单退出登录；公共或共享终端不得保存密码。", "发现异常登录、权限不符或凭证泄露时，立即停止使用并联系管理员禁用账号或轮换凭证。"])

    heading(doc, "3. 角色、工作空间与功能入口")
    para(doc, "平台采用平台级与租户级两层管理。平台管理员进入平台工作空间时管理全局基础设施、租户和平台智能体；切换到租户工作空间后，其可见功能按租户管理员口径收敛。")
    table(doc, [("平台管理员", "平台工作空间", "租户、平台用户、平台智能体、模型/引擎/MCP/系统设置"),
                ("平台管理员", "租户工作空间", "目标租户的知识库、用户授权、治理与验收"),
                ("租户管理员", "本租户", "本租户用户、租户信息、API、反馈/图谱审核、验收评测"),
                ("普通成员", "本租户", "被授权知识库、搜索、对话及个人常规设置")],
          [1900, 1800, 5660], headers=["身份", "工作空间", "主要入口"], header=False)
    note(doc, "操作前确认", "页面右上角或工作空间切换处应先确认当前租户。平台管理员在错误租户中操作，可能造成模型、用户或知识资产配置错位。", True)

    heading(doc, "4. 平台管理员操作")
    heading(doc, "4.1 租户管理", 2)
    para(doc, "路径：平台工作空间 → 租户管理。列表展示租户 ID、名称、登录用户名、存储配额、状态和创建时间。")
    bullets(doc, ["点击“新增租户”，填写租户名称与存储配额。", "可指定租户管理员登录用户名与初始密码；用户名仅支持英文字母、数字、点、下划线和短横线。", "创建后仅在受控渠道交付初始凭据，并要求首次登录立即改密。", "通过“用户”进入该租户的用户管理；通过“编辑”调整名称、配额或重置管理员密码。", "停用用于临时阻断租户访问；已投入使用的租户优先停用，不直接删除。", "只有界面明确标记为可删除的未使用租户才执行删除，并先完成业务确认。"], numbered=True)
    note(doc, "验证方法", "新建租户后使用初始管理员账号登录一次，确认租户名称、角色和可见菜单正确；停用后确认该租户账号无法继续访问。")
    table(doc, [("租户名称", "客户/部门/项目的业务标识", "建议采用统一命名规范"),
                ("存储配额", "限制租户可使用的存储容量", "结合文档量、版本和增长率设置"),
                ("登录用户名", "租户初始管理员登录名", "创建后妥善交付"),
                ("状态", "启用或停用租户访问", "停用前通知业务负责人")], [1900, 3900, 3560], headers=["字段", "作用", "管理要求"], header=False)
    heading(doc, "4.2 第三方项目接入", 2)
    para(doc, "在租户操作中可创建第三方项目接入，绑定接入项目、身份提供方、租户管理员和允许的 Origin。创建后生成的客户端密钥只应交付给受信任的服务端，浏览器端不得持有。")
    bullets(doc, ["Origin 只填写协议、域名和端口，不包含路径。", "确认接入会授权该租户当前及未来创建的知识库，必须经过数据负责人审批。", "记录客户端 ID、密钥保管人、使用系统和轮换日期。", "密钥疑似泄露时立即轮换，并验证旧密钥失效。"])
    heading(doc, "4.3 平台智能体管理", 2)
    para(doc, "路径：平台工作空间 → 智能体。平台管理员可新增、编辑、启用/停用平台智能体。配置完成后会影响租户对话中的可选能力，变更前应使用测试租户验证。")

    heading(doc, "5. 用户、角色与知识库授权")
    heading(doc, "5.1 新增用户", 2)
    para(doc, "路径：用户管理。平台管理员先在租户列表选择目标租户；租户管理员直接管理本租户用户。")
    bullets(doc, ["填写昵称、登录用户名和初始密码。用户名为 2–100 位，只能包含英文字母、数字、点、下划线和短横线。", "选择普通成员或租户管理员。平台级角色不能在租户用户页面随意授予。", "普通成员选择“全部知识库”或“指定知识库”；指定模式下逐项选择允许查看的知识库。", "创建后要求用户首次登录修改密码，并由业务负责人复核可见知识库。"] , numbered=True)
    note(doc, "验证方法", "用新增或调整后的普通成员账号登录，确认只能看到授权知识库，且无法进入租户管理、模型和基础引擎等管理员入口。")
    heading(doc, "5.2 编辑、禁用与删除", 2)
    bullets(doc, ["编辑可调整昵称、用户名、角色、知识库范围；重置密码留空表示不修改。", "租户必须至少保留一名启用的租户管理员。", "租户管理员改为普通成员时，应主动选择知识库范围，避免自动获得过宽权限。", "禁用适合离职、调岗或临时冻结；启用前重新核对角色和知识库权限。", "管理员、当前登录用户或已有文档操作记录的用户可能不允许删除，此时使用禁用保留审计链。"])
    table(doc, [("全部知识库", "当前及后续符合规则的知识库", "内部管理员或确需全量访问的岗位"),
                ("指定知识库", "只访问明确勾选的知识库", "普通业务成员，推荐默认")], [2100, 3700, 3560], headers=["授权方式", "范围", "适用场景"], header=False)

    heading(doc, "6. 平台设置与基础能力配置")
    para(doc, "平台管理员在“设置”中维护全局基础能力；租户管理员只看到租户信息、API、反馈审核、图三元组审核和验收评测等租户级入口；普通成员仅开放常规设置。")
    table(doc, [("常规设置", "界面和通用行为", "修改后用普通账号验证"),
                ("模型配置", "对话、Embedding、ReRank、视觉、ASR 等模型", "逐项连接测试并配置档位"),
                ("网络搜索", "搜索 Provider、凭证与可用性", "测试查询并控制外网访问"),
                ("向量数据库引擎", "向量检索后端连接", "连接测试后再分配给知识库"),
                ("解析引擎", "文档解析服务与策略", "用 PDF/Word/表格样例验证"),
                ("存储引擎", "本地、MinIO 或云对象存储", "上传并确认对象落盘"),
                ("MCP 服务", "外部工具、认证、超时与重试", "最小权限并查看工具清单"),
                ("系统信息", "版本与运行信息", "用于报障和变更记录")],
          [1900, 4100, 3360], headers=["设置项", "管理内容", "保存后验证"], header=False)
    heading(doc, "6.1 配置模型", 2)
    para(doc, "进入“设置 → 模型配置”，至少准备对话模型、Embedding 模型和 ReRank 模型。Embedding 用于向量化，ReRank 用于重排召回结果，对话模型负责生成回答。")
    bullets(doc, ["选择模型厂商或“自定义（OpenAI 兼容接口）”。", "填写服务地址、模型 ID 和 API Key；模型 ID 必须与服务端公开名称完全一致。", "执行连接测试，通过后保存。", "在模型 Profile 检查清单中确认在线/离线档位所需模型均已就绪。"] , numbered=True)
    note(doc, "内网模型", "如模型端点为内网地址，部署管理员需在服务端配置 SSRF 白名单或允许私网访问；业务用户不要自行绕过安全策略。")

    heading(doc, "6.2 网络搜索、引擎与 MCP", 2)
    bullets(doc, ["网络搜索：选择供应商并填写凭证，连接测试通过后才允许智能体使用；受控网络应完成外联审批。", "向量数据库：配置连接地址、认证和集合/索引信息，切换后新旧索引通常不能自动互认。", "解析引擎：按文件复杂度选择内置、Simple、MarkItDown、MinerU 或客户交付的解析服务。", "存储引擎：可选本地、MinIO、COS、TOS、S3、OSS；切换前先制定历史对象迁移方案。", "MCP 服务：支持 SSE、HTTP Streamable 或 Stdio；连接测试后检查暴露的工具清单，关闭非必要工具。"])
    note(doc, "变更控制", "模型、Embedding、向量库、解析或存储引擎属于高影响配置。生产变更前应备份、在测试知识库验证，并预留回退方案。", True)
    note(doc, "验证方法", "每项连接测试通过后，用脱敏样例完成一次上传、解析、检索和问答；涉及 MCP 或网络搜索时，再核对实际工具调用记录。")

    heading(doc, "7. 创建与配置知识库")
    bullets(doc, ["进入“知识库”，选择“新建知识库”。", "选择知识库类型：文档型适合制度、合同、报告；问答型适合标准 FAQ；Wiki 型适合自动生成结构化知识页面。", "填写名称与说明，选择解析引擎、对话模型与 Embedding 模型；启用 Wiki 时再选择 Wiki 综合模型。", "按业务需要设置分块、图像处理、音频识别、AI 问题生成和多模态。", "保存后进入知识库详情页。"], numbered=True)
    table(doc, [("文档型", "PDF、Word、PPT、Excel、Markdown、图片等", "政策制度、产品资料、项目档案"),
                ("问答型", "标准问、相似问、反例与标准答案", "客服 FAQ、办事指南"),
                ("Wiki 型", "由文档生成互相关联的 Wiki 页面", "研究资料、技术知识网络")], [1500, 3500, 4360], headers=["类型", "内容特点", "推荐场景"], header=False)

    heading(doc, "7.1 基本信息与类型", 2)
    bullets(doc, ["名称应能表达业务边界；说明中写清资料范围、责任部门和更新时间。", "文档型用于非结构化资料，问答型用于标准 FAQ，Wiki 型用于页面化知识组织。", "知识库创建后指定内容负责人和权限负责人，避免无人维护。"])
    heading(doc, "7.2 模型配置", 2)
    bullets(doc, ["对话模型用于基于知识生成回答。", "Embedding 模型用于向量索引；知识入库后通常不应直接更换，确需更换时要重建索引。", "启用 Wiki 时可单独指定 Wiki 综合模型。", "需要重排召回结果时，由平台默认配置或智能体检索配置提供 ReRank 能力。"])
    heading(doc, "7.3 索引策略、分块与高级设置", 2)
    table(doc, [("向量/关键词检索", "决定普通搜索与 RAG 召回", "用典型问题测试召回片段"),
                ("Wiki", "生成结构化页面与链接", "检查页面层级和引用"),
                ("知识图谱", "抽取实体、关系并支持多跳检索", "抽样审核三元组"),
                ("分块大小/重叠", "影响语义完整性、召回粒度和成本", "长文、表格、条款分别测试"),
                ("AI 问题生成", "为分块生成推荐问题", "控制数量和模型成本"),
                ("图像/音频", "处理图片、多模态或语音内容", "确认对应模型可用")],
          [2300, 3900, 3160], headers=["配置", "影响", "验证重点"], header=False)
    heading(doc, "7.4 存储配置", 2)
    para(doc, "知识库可在平台允许范围内选择存储引擎。变更前应确认目标存储可用，并制定历史对象迁移、校验和回退方案。")

    heading(doc, "8. 导入与维护知识")
    heading(doc, "8.1 上传文件与目录", 2)
    bullets(doc, ["在知识库详情页选择文档导入，点击或拖拽上传。", "支持 PDF、Word、TXT、Markdown、HTML、图片、CSV、Excel、PPT、JSON 等格式。", "等待解析、分块和索引完成；失败条目应查看错误信息后重试。", "打开文档检查正文、分块和图片是否正确，再开始正式问答。"], numbered=True)
    heading(doc, "8.2 URL、在线录入与外部数据源", 2)
    para(doc, "可按界面入口导入网页 URL、在线录入 Markdown/FAQ，或在“数据源管理”接入飞书、Notion、语雀等外部知识库。数据源首次接入应先验证凭证，再选择全量或增量同步，并通过同步日志确认结果。")
    heading(doc, "8.3 FAQ 管理", 2)
    para(doc, "问答型知识库支持标准问、相似问、反例和答案管理，也可批量导入。标准问用于定义意图，相似问扩大表达覆盖，反例用于降低误匹配。批量导入前先用少量样例检查列映射和字符编码。")
    heading(doc, "8.4 日常治理", 2)
    bullets(doc, ["为文档设置清晰标题、标签和文件夹；同类资料使用统一命名。", "文件更新后重新上传或触发数据源同步，避免新旧版本并存。", "删除前确认是否仍被智能体或业务接口使用。", "抽样检查解析正文、召回片段和引用来源，而不是只看“上传成功”。"])
    heading(doc, "8.5 删除与变更影响", 2)
    bullets(doc, ["删除文档会影响后续检索和引用；先确认是否有替代版本，并记录删除原因。", "删除知识库前检查智能体、API 接入和业务流程是否仍引用该知识库。", "修改 Embedding、分块或索引策略后，按界面提示重建索引并重新执行验收问题集。", "删除用户优先改为禁用以保留审计链；删除智能体前先确认没有生产对话依赖。"])

    heading(doc, "9. 搜索与智能问答")
    heading(doc, "9.1 知识搜索", 2)
    bullets(doc, ["进入“搜索”，输入能够表达业务意图的关键词或问题。", "按知识库、文档或标签缩小范围。", "打开命中片段核对来源；无结果时换用同义词、缩短问题或检查文档索引状态。"], numbered=True)
    heading(doc, "9.2 发起对话", 2)
    bullets(doc, ["进入“对话/新建对话”，选择合适智能体。", "通过 @ 选择知识库或文件；智能体配置可能限制可选范围。", "输入明确问题，必要时说明时间范围、对象、输出格式和引用要求。", "阅读回答时展开引用，核对原文；关键结论应回到源文件确认。", "继续追问时补充约束，不要假设模型自动知道客户内部未入库的信息。"], numbered=True)
    table(doc, [("快速问答", "单次知识检索与回答", "事实查询、制度定位"),
                ("智能推理", "多步推理并可调用工具", "跨文档分析、复杂任务"),
                ("深度研究员", "研究型检索与综合", "专题调研、证据汇总"),
                ("数据分析师", "处理数据与文件", "表格分析、结构化结果"),
                ("知识图谱专家", "结合实体关系理解知识", "关系追踪、多跳问题")], [1900, 3300, 4160], headers=["智能体", "特点", "适用任务"], header=False)

    heading(doc, "10. Wiki 与知识图谱")
    para(doc, "Wiki 模式可从原始文档生成结构化 Markdown 页面及页面间链接。启用图谱后，系统可抽取实体与关系，并在检索时补充上下文。")
    bullets(doc, ["确认部署侧已启用 Neo4j 和 GraphRAG。", "在知识库设置中开启实体/关系抽取，并按业务定义 Schema、实体数、关系数和置信度。", "上传或重新处理文档，等待图谱任务完成。", "在 Wiki 浏览器或图谱视图检查页面、节点与关系；错误三元组应按审核流程处理。"] , numbered=True)
    note(doc, "成本与时效", "图谱构建需要调用大模型，耗时和费用通常高于普通文档索引。只在多跳关系确有价值的知识库中启用。")

    heading(doc, "11. 智能体配置与发布")
    para(doc, "平台管理员可创建或维护自定义智能体，配置模型、系统 Prompt、知识库范围、检索策略、最大迭代次数、网络搜索、MCP 服务、Skills 以及图片/语音/附件能力。")
    bullets(doc, ["先定义智能体任务边界，再选择普通模式或智能推理模式。", "知识库选择“全部/指定/不使用”时，确认与启用工具兼容。", "仅开放完成任务必需的 MCP/Skills，并使用最小权限凭证。", "保存前确认对话模型与 ReRank 模型就绪；用典型问题和失败问题各测试一次。"])
    heading(doc, "11.1 基础与模型配置", 2)
    table(doc, [("名称/头像/描述", "决定用户识别和选择", "名称体现业务用途"),
                ("智能体类型/模式", "普通对话或多步智能推理", "复杂任务才使用 Agent 模式"),
                ("思考/对话模型", "影响能力、速度与成本", "使用批准模型并设合理超时"),
                ("最大迭代次数", "限制工具调用循环", "过高会增加费用和等待"),
                ("温度/Token", "影响随机性与输出长度", "知识问答宜较稳定")],
          [2200, 3900, 3260], headers=["配置项", "影响", "建议"], header=False)
    heading(doc, "11.2 知识库、检索与工具", 2)
    bullets(doc, ["知识库范围可选全部、指定或不使用；指定模式应按最小权限选择。", "普通模式可配置历史保留轮数、Embedding TopK、关键词/向量阈值、ReRank TopK/阈值、问题改写、查询扩展和兜底策略。", "Agent 模式按任务开放知识检索、Wiki、网络搜索、MCP 和 Skills；工具越多，误调用面和审核成本越高。", "图片、语音、附件和 @ 按钮应与模型能力、文件类型及业务场景一致。", "系统 Prompt 写清角色、可用数据、禁止事项、引用要求和失败处理，不写入真实密钥。"])
    heading(doc, "11.3 测试与启停", 2)
    bullets(doc, ["准备正常问题、边界问题、无答案问题、越权问题和恶意提示词进行测试。", "检查回答、引用、工具调用、耗时和失败提示，不只看是否有文字输出。", "配置不完整的智能体不要启用；停用后确认对话下拉不再出现。"])

    heading(doc, "12. 租户治理、API 与验收评测")
    heading(doc, "12.1 租户信息与 API 信息", 2)
    para(doc, "租户管理员可查看和维护租户信息，并在 API 信息中获取本租户集成所需信息。API 凭证按服务账号管理，禁止嵌入前端页面或公开脚本；停用人员不应继续持有凭证。")
    heading(doc, "12.2 回答反馈与图三元组审核", 2)
    bullets(doc, ["定期查看用户反馈，按知识缺失、召回错误、生成错误或权限问题分类。", "图三元组审核应核对实体、关系、来源片段和业务 Schema；错误关系拒绝或修正后再写入图谱。", "审核结论应能追溯到知识来源和处理人。"])
    heading(doc, "12.3 验收评测", 2)
    para(doc, "使用代表性问题集验证检索与生成效果。至少记录问题、期望来源、实际召回、回答结论、引用、是否通过和失败原因；配置变更或模型升级后重复执行同一基准。")

    heading(doc, "13. 常见问题速查")
    table(doc, [("无法登录", "核对地址、账号状态、租户与密码；联系管理员重置"),
                ("模型连接失败", "核对模型 ID、服务地址、API Key、网络与 SSRF 策略"),
                ("上传后一直处理中", "检查文件大小/格式，联系运维查看 app、docreader、队列日志"),
                ("搜索不到内容", "确认解析和索引成功，检查知识库权限、阈值和模型配置"),
                ("回答没有引用", "确认选中了知识库且智能体允许知识检索；检查召回是否为空"),
                ("图片链接无效", "确认多模态配置、MinIO 服务与公开访问地址"),
                ("配置保存后消失", "刷新并复核接口结果；联系运维检查后端日志和租户上下文")], [2300, 7060], headers=["现象", "处理建议"])

    heading(doc, "14. 分角色培训验收清单")
    heading(doc, "14.1 平台管理员", 2)
    bullets(doc, ["能够新增、编辑、停用租户并解释删除限制。", "能够进入指定租户管理用户与知识库范围。", "能够完成模型、搜索、解析、向量、存储和 MCP 的连接测试。", "能够配置并测试平台智能体，理解启停影响。"])
    heading(doc, "14.2 租户管理员", 2)
    bullets(doc, ["能够新增/禁用用户并按最小权限授权知识库。", "能够创建并完整配置文档型、问答型或 Wiki 知识库。", "能够管理 FAQ、数据源、反馈和图三元组审核。", "能够使用固定问题集完成一次验收评测。"])
    heading(doc, "14.3 普通成员", 2)
    bullets(doc, ["能够登录并识别当前角色/租户。", "能够检索授权知识库并打开来源片段。", "能够用 @ 指定知识库发起问答并核对引用。", "能够选择合适智能体并识别无答案或越权请求。"])
    note(doc, "培训建议", "现场演示使用脱敏样例文档。每位学员至少完成一次“上传 → 检索 → 问答 → 核对引用”的闭环。")
    path = OUT / "企业知识智能平台用户使用手册（客户培训版）.docx"
    doc.save(path)
    return path


def deployment_manual():
    doc = Document()
    setup(doc, "部署手册")
    cover(doc, "部署手册", "生产环境 Docker Compose · 安装、验收与运维", "实施工程师、系统管理员、运维与安全人员")
    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_before = Pt(0)
    toc_title.paragraph_format.space_after = Pt(12)
    font(toc_title.add_run("目录"), 20, True, NAVY)
    toc_placeholder = doc.add_paragraph("[[TOC]]")
    toc_placeholder.paragraph_format.space_after = Pt(0)
    doc.add_page_break()
    heading(doc, "1. 部署目标与边界")
    para(doc, "本手册用于把企业知识智能平台部署到远程 Linux 服务器。默认使用客户交付的实施包构建或导入镜像，由 Docker Compose 管理应用、数据库、检索、对象存储和可选服务。Kubernetes 或严格离线环境应使用对应的专项交付包与实施方案。")
    note(doc, "默认假设", "安装目录 /opt/knowledge-platform；客户通过反向代理或指定前端端口访问；生产密钥由客户的 Secret 管理流程生成和保管。")

    heading(doc, "2. 架构与服务组成")
    para(doc, "核心链路为：浏览器 → Web UI → 应用服务 → PostgreSQL/Redis → 文档解析服务 → Elasticsearch/Milvus → 对象存储；按需扩展 Neo4j、Qdrant、Langfuse、Jaeger、Dex 和 Agent Sandbox。")
    table(doc, [("核心", "app、frontend、docreader、postgres、redis、elasticsearch、milvus", "登录、解析、存储、检索与问答"),
                ("full 扩展", "minio、qdrant、neo4j、jaeger、dex、langfuse、sandbox", "对象存储、图谱、可观测、OIDC、技能执行"),
                ("不含于 full", "weaviate", "需要时显式启用 full,weaviate")], [1800, 3800, 3760], headers=["分组", "组件", "说明"], header=False)
    heading(doc, "2.1 部署模式选择", 2)
    table(doc, [("基础生产", "不设置 Profile", "核心问答；本地存储或已接外部存储"),
                ("标准生产", "minio", "核心问答 + MinIO 对象存储，推荐常规私有化项目"),
                ("图谱增强", "minio,neo4j", "标准生产 + 知识图谱"),
                ("全功能验证", "full", "包含观测、OIDC、图谱等多数可选组件；不建议无差别用于生产"),
                ("专项扩展", "按需叠加 qdrant / langfuse / jaeger / weaviate", "仅在需求与验收范围明确时启用")],
          [1800, 3000, 4560], headers=["模式", "COMPOSE_PROFILES", "适用范围"], header=False)
    note(doc, "选择原则", "先确定客户实际使用的存储、检索、图谱、认证和可观测能力，再启动相应 Profile；容器启动越多，资源、暴露面和运维成本越高。")

    heading(doc, "3. 资源、软件与网络准备")
    table(doc, [("CPU", "至少 4 核", "完整服务及高并发建议 8 核以上"), ("内存", "至少 16 GB", "需计入向量库、图谱和模型服务"), ("磁盘", "至少 40 GB 可用", "另计文档、索引、数据库、镜像、日志增长"), ("软件", "Docker、Docker Compose、Git", "由客户批准的稳定版本")], [1600, 2600, 5160], headers=["项目", "最低建议", "说明"], header=False)
    para(doc, "上线前执行：")
    code_lines = ["docker --version", "docker compose version", "df -h /", "free -h", "sudo ss -lntup"]
    code(doc, code_lines)
    bullets(doc, ["确认服务器时区与客户要求一致，并使用 chrony/ntpd 保持时间同步。", "确认业务域名、DNS 解析、TLS 证书及证书续期责任人。", "如需访问外部模型、对象存储或搜索服务，确认代理、出口白名单和证书信任链。", "记录服务器主机名、内网 IP、操作系统和 Docker 版本。"])
    table(doc, [("Web UI", "80", "FRONTEND_PORT"), ("App API", "8080", "APP_PORT"), ("DocReader", "50051", "DOCREADER_PORT"), ("PostgreSQL", "127.0.0.1:5432", "DB_PORT"), ("Redis", "127.0.0.1:6379", "REDIS_PORT"), ("Elasticsearch", "127.0.0.1:9200", "ELASTICSEARCH_PORT"), ("MinIO", "9000 / 9001", "MINIO_PORT / MINIO_CONSOLE_PORT"), ("Qdrant", "6333 / 6334", "QDRANT_REST_PORT / QDRANT_PORT"), ("Langfuse", "3000", "LANGFUSE_WEB_PORT"), ("Neo4j", "7474 / 7687", "Compose 当前固定"), ("Jaeger", "16686", "Compose 当前固定")], [2200, 2500, 4660], headers=["服务", "默认绑定/端口", "配置项"], header=False)
    note(doc, "防火墙", "数据库、Redis、Elasticsearch 默认仅绑定 127.0.0.1。生产环境通常只应向客户网段开放 HTTPS/网关端口；其余管理端口按最小暴露原则限制。", True)
    table(doc, [("业务用户网段", "443（或客户统一网关端口）", "Web UI 与 API"),
                ("运维管理网段", "SSH、必要的 MinIO/Langfuse/Neo4j/Jaeger 管理端口", "按启用组件逐项放行"),
                ("宿主机本地", "PostgreSQL、Redis、Elasticsearch", "保持 127.0.0.1 绑定"),
                ("容器网络", "服务间内部端口", "不映射到公网")], [2100, 3400, 3860], headers=["来源", "允许端口", "要求"], header=False)

    heading(doc, "4. 获取代码与建立配置")
    code(doc, ["sudo mkdir -p /opt/knowledge-platform", "sudo chown \"$(id -u):$(id -g)\" /opt/knowledge-platform", "cd /opt/knowledge-platform", "# 将已验签的客户交付包解压到当前目录", "cp .env.example .env", "chmod 600 .env"])
    para(doc, "交付包应通过客户批准的介质传输并完成校验。不要上传本地 .env、版本控制目录、依赖缓存或无关构建文件。")
    code(doc, ["sha256sum <交付包文件>", "docker compose images --quiet | sort", "# 将交付包版本、SHA-256 和镜像摘要写入交付记录"])

    heading(doc, "5. 生产配置")
    heading(doc, "5.1 必改项", 2)
    code(doc, ["GIN_MODE=release", "COMPOSE_PROFILES=minio", "FRONTEND_PORT=8089", "APP_PORT=18080", "DOCREADER_PORT=15051", "DB_PORT=15432", "REDIS_PORT=16379", "", "DEFAULT_ADMIN_ENABLED=true", "DEFAULT_ADMIN_USERNAME=admin", "DEFAULT_ADMIN_EMAIL=admin@example.com", "DEFAULT_ADMIN_PASSWORD=<随机强密码>", "", "DB_PASSWORD=<随机强密码>", "REDIS_PASSWORD=<随机强密码>", "ELASTICSEARCH_PASSWORD=<随机强密码>", "JWT_SECRET=<随机值>", "TENANT_AES_KEY=<随机值>", "SYSTEM_AES_KEY=<随机值>", "# OIDC/Dex 客户端密钥变量请按交付包 .env.example 填写"])
    para(doc, "随机值可使用 openssl rand -hex 32 或 openssl rand -base64 32 生成。SYSTEM_AES_KEY 等有长度要求时，以 .env.example 注释为准。")
    heading(doc, "5.2 存储与检索驱动", 2)
    code(doc, ["STORAGE_TYPE=minio", "RETRIEVE_DRIVER=elasticsearch_v8,milvus", "KEYWORD_RETRIEVE_DRIVER=elasticsearch_v8", "VECTOR_RETRIEVE_DRIVER=milvus", "ENABLE_GRAPH_RAG=false", "NEO4J_ENABLE=false", "# 启用图谱时再将 Profile 加入 neo4j，并把上述两个开关改为 true"])
    note(doc, "关键概念", "Profile 决定容器是否启动；STORAGE_TYPE 和各 RETRIEVE_DRIVER 决定 App 实际使用哪个后端。两者必须一致。")
    heading(doc, "5.3 生产安全基线", 2)
    bullets(doc, ["所有默认密码、JWT/AES/Dex/Langfuse 密钥必须替换；.env 权限保持 600。", "真实凭证不得进入 Git、工单、聊天或构建日志。", "模型、对象存储、MCP 使用独立最小权限账号并定期轮换。", "SSRF_WHITELIST 只加入确需访问的域名/IP/CIDR；不得用大网段图省事。", "通过 HTTPS 反向代理提供 Web 服务，并限制管理端口来源。", "建立数据库、对象存储和配置的备份策略；定期演练恢复。"])
    heading(doc, "5.4 HTTPS 反向代理参考", 2)
    code(doc, ["server {", "  listen 443 ssl http2;", "  server_name knowledge.example.com;", "  ssl_certificate /etc/nginx/certs/fullchain.pem;", "  ssl_certificate_key /etc/nginx/certs/privkey.pem;", "  client_max_body_size 2047m;", "  location / {", "    proxy_pass http://127.0.0.1:8089;", "    proxy_http_version 1.1;", "    proxy_set_header Host $host;", "    proxy_set_header X-Real-IP $remote_addr;", "    proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;", "    proxy_set_header X-Forwarded-Proto https;", "    proxy_read_timeout 600s;", "    proxy_send_timeout 600s;", "    proxy_buffering off;", "  }", "}"])
    note(doc, "网关注意", "示例端口与路径需按客户网关调整。流式回答依赖较长超时和关闭代理缓冲；上传限制应与 MAX_FILE_SIZE_MB 保持一致。证书私钥只允许受权运维人员读取。")

    heading(doc, "6. 配置校验、构建与启动")
    para(doc, "先解析 Compose；任何错误都应先修复：")
    code(doc, ["cd /opt/knowledge-platform", "docker compose config --services", "docker compose config --quiet"])
    para(doc, "构建并启动：")
    code(doc, ["docker compose build --pull app frontend docreader sandbox", "docker compose up -d", "docker compose ps"])
    note(doc, "首次启动", "PostgreSQL、Elasticsearch、Milvus、ClickHouse 和 Langfuse 初始化可能持续数分钟。sandbox 准备容器执行 true 后退出属于正常行为。")

    heading(doc, "7. 健康检查与交付验收")
    code(doc, ["docker compose ps", "docker compose logs --tail=100 app", "docker compose logs --tail=100 langfuse-web", "curl --fail http://127.0.0.1:18080/health", "curl --head http://127.0.0.1:8089"])
    bullets(doc, ["常驻服务均为 running 或 healthy。", "App /health 成功，Web UI 可打开并登录。", "立即修改默认管理员密码。", "配置一组可用的对话、Embedding、ReRank 模型并通过连接测试。", "上传脱敏测试文档，确认解析、索引、搜索、问答和引用全链路成功。", "确认对象进入目标存储，检索引擎无连接错误。", "启用 Langfuse 时，问答后能看到 Trace。", "启用 Neo4j 时，图谱抽取和查询通过。"])
    table(doc, [("系统地址", "____________________________"), ("管理员账号", "____________________________"), ("版本/提交", "____________________________"), ("部署日期", "____________________________"), ("实施/客户确认", "____________________________")], [2300, 7060], headers=["交付项", "现场记录"])

    heading(doc, "8. Langfuse 初始化（可选）")
    bullets(doc, ["打开 http://服务器地址:3000。", "创建管理员、组织和项目。", "在 Settings → API Keys 生成 Public Key 和 Secret Key。", "写入 LANGFUSE_HOST=http://langfuse-web:3000、Public/Secret Key。", "执行 docker compose up -d --force-recreate app，并发起一次问答验证。"] , numbered=True)
    note(doc, "调试", "LANGFUSE_DEBUG=true 仅在排障期间临时使用，问题解决后恢复 false。")

    heading(doc, "9. 日常运维")
    table(doc, [("查看状态", "docker compose ps"), ("跟踪应用日志", "docker compose logs -f --tail=200 app"), ("停止/启动", "docker compose stop / docker compose start"), ("重建应用", "docker compose up -d --force-recreate app"), ("停止并移除容器网络", "docker compose down（默认保留命名卷）")], [2600, 6760], headers=["操作", "命令"])
    note(doc, "禁止误操作", "除非明确要永久删除全部持久化数据，不要执行 docker compose down -v。共享服务器不要直接执行 docker system prune -a。", True)
    heading(doc, "9.1 监控建议", 2)
    bullets(doc, ["CPU、内存、磁盘/ inode、容器重启次数。", "App 健康检查、请求错误率、P95 延迟。", "PostgreSQL 连接数与备份结果；Redis/队列积压。", "Elasticsearch/Milvus/Qdrant 容量与错误。", "对象存储容量和失败上传；Langfuse/Jaeger 链路。", "模型 API 可用性、限流、Token/费用与平均耗时。"])
    heading(doc, "9.2 日志与容量治理", 2)
    code(doc, ["# /etc/docker/daemon.json 示例", "{", "  \"log-driver\": \"json-file\",", "  \"log-opts\": { \"max-size\": \"100m\", \"max-file\": \"5\" }", "}"])
    bullets(doc, ["变更 Docker 日志配置后在维护窗口重启 Docker，并验证全部容器恢复。", "磁盘使用率达到 70% 预警、85% 严重告警；同时监控 inode。", "对象、数据库、索引、备份和容器日志分别设置保留策略，禁止在未确认数据归属前直接清理卷。"])

    heading(doc, "10. 备份、升级与回滚")
    heading(doc, "10.1 升级前备份", 2)
    code(doc, ["cd /opt/knowledge-platform", "# 记录交付包版本号与镜像摘要", "docker compose images", "mkdir -p backups", "docker compose exec -T postgres sh -c 'pg_dump -U \"$POSTGRES_USER\" \"$POSTGRES_DB\"' > \"backups/postgres-$(date +%F-%H%M%S).sql\""])
    para(doc, "同时备份 .env（加密保存）、对象存储数据和外部向量/图数据库。仅备份 PostgreSQL 不代表完整业务可恢复。")
    code(doc, ["# MinIO 示例：先在受控备份机配置别名", "mc mirror --overwrite production/knowledge /backup/knowledge", "# 加密备份配置", "tar -czf - .env docker-compose.yml | openssl enc -aes-256-cbc -salt -out backups/config-$(date +%F).tgz.enc"])
    heading(doc, "10.2 恢复与演练", 2)
    code(doc, ["# 维护窗口内停止应用写入", "docker compose stop app frontend", "# 恢复 PostgreSQL（目标库应按演练方案准备）", "cat backups/postgres-<时间>.sql | docker compose exec -T postgres psql -U \"$POSTGRES_USER\" \"$POSTGRES_DB\"", "# 恢复对象存储示例", "mc mirror --overwrite /backup/knowledge production/knowledge", "docker compose start app frontend"])
    bullets(doc, ["向量索引未备份时，按知识库执行重新索引；图谱启用时按所用数据库工具恢复或重建。", "恢复后验证登录、文档列表、对象下载、检索、问答和引用，不以容器启动作为成功标准。", "明确 RPO、RTO、备份频率、保留周期、备份加密和恢复责任人；至少每季度演练一次。"])
    heading(doc, "10.3 升级", 2)
    code(doc, ["# 将目标版本交付包部署到受控目录", "docker compose config --quiet", "docker compose build --pull app frontend docreader sandbox", "docker compose up -d"])
    para(doc, "升级后完整执行第 7 节验收，尤其是上传、解析、检索和问答闭环。")
    heading(doc, "10.4 回滚", 2)
    para(doc, "应用代码可切回升级前提交后重新构建。数据库迁移不保证可逆；若发生不兼容迁移，应停止写入，按已演练方案恢复数据库及关联存储。容器启动不等于回滚成功。")
    note(doc, "变更流程", "升级前完成审批、通知、停止写入、备份和版本记录；升级后执行健康检查与端到端验收。任一关键验收失败且无法在窗口内修复时，立即按预案回滚。")

    heading(doc, "11. 故障排查")
    table(doc, [("Compose 报必填变量为空", "补齐 .env，重新执行 docker compose config --quiet"),
                ("端口/容器冲突", "用 docker ps -a、ss -lntup 定位；调整变量或明确停用冲突服务"),
                ("Web 可开但 API 失败", "检查 app 健康、前端代理目标、端口与防火墙"),
                ("文档上传/解析失败", "检查 MAX_FILE_SIZE_MB、app/docreader/Redis 队列日志及对象存储"),
                ("MinIO/Qdrant 已启动但未使用", "检查 STORAGE_TYPE 与 RETRIEVE_DRIVER，而不只看 Profile"),
                ("Langfuse 无 Trace", "容器内 Host 应为 http://langfuse-web:3000；核对项目 Key 并重建 app"),
                ("模型访问被拒", "核对网络、TLS、模型 ID、凭证和 SSRF 配置"),
                ("磁盘快速增长", "定位卷、日志、对象与索引；按保留策略清理，先备份再操作")], [2600, 6760], headers=["现象", "处理路径"])

    heading(doc, "12. 上线检查表")
    bullets(doc, ["资源、NTP、DNS、证书、端口和防火墙已确认。", "交付包 SHA-256、镜像摘要和实际部署模式已记录。", "所有默认密钥已替换，.env 已受控保存。", "Compose 配置解析通过，镜像版本已记录。", "健康检查和端到端业务验收通过。", "管理员密码已修改，角色和知识库授权符合最小权限。", "备份任务已配置，并完成一次恢复演练。", "监控告警、日志保留、应急联系人及升级回滚责任人已登记。"])
    path = OUT / "企业知识智能平台部署手册（客户培训版）.docx"
    doc.save(path)
    return path


def code(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [9360])
    shade(t.cell(0, 0), "F2F4F7")
    p = t.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        font(p.add_run(line or " "), 8.8, color="1F2937", name="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


if __name__ == "__main__":
    print(user_manual())
    print(deployment_manual())
